"""
Document parser for CSES variable mapping.

Handles: .docx, .pdf, .txt, .rtf, .doc
Extracts: text content, questionnaire structure (Q01, Q02, D01, etc.)
"""

import logging
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class QuestionInfo:
    """Information about a questionnaire item."""
    code: str  # Q01, Q02a, D04, etc.
    text: str  # The question text
    response_options: list[str] = field(default_factory=list)


@dataclass
class DocumentInfo:
    """Information about a parsed document."""
    file_path: Path
    file_format: str
    full_text: str
    questions: list[QuestionInfo]
    is_questionnaire: bool
    language_detected: Optional[str] = None
    parse_errors: list[str] = field(default_factory=list)

    def summary(self) -> str:
        """Get a summary of the document."""
        return (
            f"File: {self.file_path.name}\n"
            f"Format: {self.file_format}\n"
            f"Text length: {len(self.full_text)} chars\n"
            f"Questions found: {len(self.questions)}\n"
            f"Is questionnaire: {self.is_questionnaire}"
        )


class DocumentParser:
    """Format-agnostic document parser."""

    SUPPORTED_FORMATS = {
        '.docx': 'docx',
        '.doc': 'doc',
        '.pdf': 'pdf',
        '.txt': 'txt',
        '.rtf': 'rtf',
        '.md': 'txt',
    }

    # Patterns for questionnaire items (case-insensitive)
    # NOTE: These are MINIMAL patterns - the real work is done by the LLM
    # Collaborator documents are messy and diverse - don't assume specific formats
    QUESTION_PATTERNS = [
        # Q01, Q02a, Q02b format
        r'(?P<code>Q\d{1,2}[a-z]?)[.\s:)\]]+\s*(?P<text>[^\n]+)',
        # D01, D02, D04 format (demographics)
        r'(?P<code>D\d{1,2}[a-z]?)[.\s:)\]]+\s*(?P<text>[^\n]+)',
        # A01, A02 format (admin)
        r'(?P<code>A\d{1,2}[a-z]?)[.\s:)\]]+\s*(?P<text>[^\n]+)',
        # Question 1, Question 2 format
        r'(?:Question|Pregunta|Frage|Domanda)\s*(?P<code>\d{1,2})[.\s:)\]]+\s*(?P<text>[^\n]+)',
    ]

    def __init__(self):
        self.errors = []

    def parse(self, file_path: Path | str) -> Optional[DocumentInfo]:
        """
        Parse a document and extract content.

        Args:
            file_path: Path to the document

        Returns:
            DocumentInfo object or None if parsing fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            logger.error(f"Unsupported format: {suffix}")
            return None

        format_type = self.SUPPORTED_FORMATS[suffix]

        try:
            if format_type == 'docx':
                return self._parse_docx(file_path)
            elif format_type == 'doc':
                return self._parse_doc(file_path)
            elif format_type == 'pdf':
                return self._parse_pdf(file_path)
            elif format_type == 'txt':
                return self._parse_txt(file_path)
            elif format_type == 'rtf':
                return self._parse_rtf(file_path)
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            # Return partial result with error
            return DocumentInfo(
                file_path=file_path,
                file_format=format_type,
                full_text="",
                questions=[],
                is_questionnaire=False,
                parse_errors=[str(e)]
            )

    def _parse_docx(self, file_path: Path) -> DocumentInfo:
        """Parse DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(str(file_path))

        # Extract all text
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        full_text = "\n".join(paragraphs)
        questions = self._extract_questions(full_text)

        return DocumentInfo(
            file_path=file_path,
            file_format="Word Document (.docx)",
            full_text=full_text,
            questions=questions,
            is_questionnaire=len(questions) >= 5  # Heuristic
        )

    def _parse_doc(self, file_path: Path) -> DocumentInfo:
        """Parse legacy DOC file."""
        # Try using antiword or textract
        import subprocess

        try:
            # Try antiword first
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                full_text = result.stdout
            else:
                # Fallback: try to read as binary and extract text
                full_text = self._extract_text_from_binary(file_path)
        except FileNotFoundError:
            # antiword not installed, try binary extraction
            full_text = self._extract_text_from_binary(file_path)

        questions = self._extract_questions(full_text)

        return DocumentInfo(
            file_path=file_path,
            file_format="Word Document (.doc)",
            full_text=full_text,
            questions=questions,
            is_questionnaire=len(questions) >= 5
        )

    def _parse_pdf(self, file_path: Path) -> DocumentInfo:
        """Parse PDF file."""
        try:
            import pypdf
        except ImportError:
            try:
                import PyPDF2 as pypdf
            except ImportError:
                raise ImportError("pypdf not installed. Run: pip install pypdf")

        reader = pypdf.PdfReader(str(file_path))

        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

        full_text = "\n".join(pages_text)
        questions = self._extract_questions(full_text)

        return DocumentInfo(
            file_path=file_path,
            file_format="PDF",
            full_text=full_text,
            questions=questions,
            is_questionnaire=len(questions) >= 5
        )

    def _parse_txt(self, file_path: Path) -> DocumentInfo:
        """Parse plain text file."""
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                full_text = file_path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            full_text = file_path.read_text(encoding='utf-8', errors='replace')

        questions = self._extract_questions(full_text)

        return DocumentInfo(
            file_path=file_path,
            file_format="Plain Text",
            full_text=full_text,
            questions=questions,
            is_questionnaire=len(questions) >= 5
        )

    def _parse_rtf(self, file_path: Path) -> DocumentInfo:
        """Parse RTF file."""
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError("striprtf not installed. Run: pip install striprtf")

        rtf_content = file_path.read_text(errors='replace')
        full_text = rtf_to_text(rtf_content)

        questions = self._extract_questions(full_text)

        return DocumentInfo(
            file_path=file_path,
            file_format="Rich Text Format (.rtf)",
            full_text=full_text,
            questions=questions,
            is_questionnaire=len(questions) >= 5
        )

    def _extract_text_from_binary(self, file_path: Path) -> str:
        """Fallback: extract printable text from binary file."""
        content = file_path.read_bytes()
        # Extract ASCII/UTF-8 printable sequences
        text_chunks = []
        current_chunk = []

        for byte in content:
            if 32 <= byte <= 126 or byte in (9, 10, 13):  # Printable or whitespace
                current_chunk.append(chr(byte))
            else:
                if len(current_chunk) > 10:  # Only keep chunks > 10 chars
                    text_chunks.append(''.join(current_chunk))
                current_chunk = []

        if current_chunk:
            text_chunks.append(''.join(current_chunk))

        return '\n'.join(text_chunks)

    def _extract_questions(self, text: str) -> list[QuestionInfo]:
        """
        Extract questionnaire items from text.

        NOTE: This is a BEST-EFFORT extraction. Collaborator documents are messy
        and diverse. The real understanding happens when the LLM processes the
        full document text. Don't rely heavily on this structured extraction.
        """
        questions = []
        seen_codes = set()

        for pattern in self.QUESTION_PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE):
                code = match.group('code').upper()
                question_text = match.group('text').strip()

                # Normalize code (Q1 -> Q01)
                code = self._normalize_code(code)

                if code not in seen_codes and len(question_text) > 10:
                    questions.append(QuestionInfo(
                        code=code,
                        text=question_text[:200]  # Truncate long questions
                    ))
                    seen_codes.add(code)

        # Sort by code
        questions.sort(key=lambda q: (q.code[0], int(re.search(r'\d+', q.code).group()) if re.search(r'\d+', q.code) else 0))

        return questions

    def _normalize_code(self, code: str) -> str:
        """Normalize question codes (Q1 -> Q01, D4 -> D04)."""
        match = re.match(r'([A-Z])(\d+)([a-z]?)', code, re.IGNORECASE)
        if match:
            prefix = match.group(1).upper()
            num = int(match.group(2))
            suffix = match.group(3).lower() if match.group(3) else ''
            return f"{prefix}{num:02d}{suffix}"
        return code.upper()


def parse_document(file_path: Path | str) -> Optional[DocumentInfo]:
    """Convenience function to parse a document."""
    parser = DocumentParser()
    return parser.parse(file_path)
