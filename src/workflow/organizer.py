"""
File Organizer for CSES Processing.

Detects collaborator files and organizes them into a single clean folder
with all files well-named using the country_year prefix.
"""

import logging
import shutil
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


# File type detection patterns
DATA_EXTENSIONS = {".dta", ".sav", ".csv", ".xlsx", ".xls", ".json", ".parquet"}
DOC_EXTENSIONS = {".docx", ".doc", ".pdf", ".txt", ".rtf"}

# Keywords for file type identification
QUESTIONNAIRE_KEYWORDS = [
    "questionnaire", "fragebogen", "cuestionario", "questionario",
    "survey", "translation", "pregunta", "enquete", "encuesta"
]
CODEBOOK_KEYWORDS = [
    "codebook", "code book", "variable", "dictionary", "codebuch",
    "diccionario", "dicionario"
]
DESIGN_KEYWORDS = [
    "design", "methodology", "technical", "rapport", "bericht", "informe"
]
MACRO_REPORT_KEYWORDS = [
    "macroreport", "macro report", "macro_report", "macro-report",
    "macro data", "macro_data", "macrodata"
]

# Keywords for district data (election results by constituency)
DISTRICT_DATA_KEYWORDS = [
    "district", "constituency", "riding", "wahlkreis", "circonscription"
]

# Keywords to identify English translation questionnaires
ENGLISH_KEYWORDS = [
    "english", "eng", "_en_", "_en.", "(en)", "[en]",
    "translation", "translated", "engl"
]

# Preferred format order for documents (first = most preferred)
DOC_FORMAT_PREFERENCE = [".pdf", ".docx", ".doc", ".rtf", ".txt"]

# Preferred format order for data files (first = most preferred)
DATA_FORMAT_PREFERENCE = [".dta", ".sav", ".parquet", ".csv", ".xlsx", ".xls", ".json"]


def filename_similarity(name1: str, name2: str) -> float:
    """
    Calculate similarity between two filenames (without extensions).
    Uses Jaccard similarity on character trigrams.

    Returns:
        Float between 0 and 1, where 1 = identical
    """
    def get_trigrams(s: str) -> set:
        s = s.lower()
        if len(s) < 3:
            return {s}
        return {s[i:i+3] for i in range(len(s) - 2)}

    trigrams1 = get_trigrams(name1)
    trigrams2 = get_trigrams(name2)

    if not trigrams1 or not trigrams2:
        return 0.0

    intersection = len(trigrams1 & trigrams2)
    union = len(trigrams1 | trigrams2)

    return intersection / union if union > 0 else 0.0


def deduplicate_by_format(
    files: list[Path],
    format_preference: list[str] = None,
    similarity_threshold: float = 0.7
) -> list[Path]:
    """
    Remove duplicate files that exist in multiple formats.
    Keeps the preferred format based on the preference list.

    Only deduplicates if filenames are similar enough (above threshold).
    All originals are still preserved in original_deposit/.

    Args:
        files: List of file paths
        format_preference: Ordered list of preferred extensions (first = most preferred)
                          Defaults to DOC_FORMAT_PREFERENCE
        similarity_threshold: Minimum similarity to consider files as duplicates

    Returns:
        Deduplicated list with preferred formats
    """
    if len(files) <= 1:
        return files

    if format_preference is None:
        format_preference = DOC_FORMAT_PREFERENCE

    # Group files by their base name similarity
    groups = []  # List of lists of similar files
    used = set()

    for i, file1 in enumerate(files):
        if i in used:
            continue

        group = [file1]
        used.add(i)
        stem1 = file1.stem

        for j, file2 in enumerate(files):
            if j in used:
                continue
            stem2 = file2.stem

            if filename_similarity(stem1, stem2) >= similarity_threshold:
                group.append(file2)
                used.add(j)

        groups.append(group)

    # For each group, pick the preferred format
    result = []
    for group in groups:
        if len(group) == 1:
            result.append(group[0])
        else:
            # Sort by format preference
            def format_rank(f: Path) -> int:
                ext = f.suffix.lower()
                try:
                    return format_preference.index(ext)
                except ValueError:
                    return len(format_preference)

            group.sort(key=format_rank)
            preferred = group[0]

            # Log what we're deduplicating
            others = [f.name for f in group[1:]]
            logger.info(f"Deduplicating: keeping {preferred.name}, skipping {others}")

            result.append(preferred)

    return result


@dataclass
class DetectedFiles:
    """Results of file detection in a directory."""
    data_files: list[Path] = field(default_factory=list)  # Survey data (main)
    district_data_files: list[Path] = field(default_factory=list)  # District/constituency election results
    questionnaire_files: list[Path] = field(default_factory=list)
    codebook_files: list[Path] = field(default_factory=list)
    design_report_files: list[Path] = field(default_factory=list)
    macro_report_files: list[Path] = field(default_factory=list)
    other_files: list[Path] = field(default_factory=list)

    # Detected country/year (if identifiable)
    country: Optional[str] = None
    country_code: Optional[str] = None
    year: Optional[str] = None

    def has_minimum_requirements(self) -> bool:
        """Check if minimum required files are present."""
        return len(self.data_files) > 0 and len(self.questionnaire_files) > 0

    def summary(self) -> str:
        """Generate summary of detected files."""
        lines = ["## Detected Files"]

        if self.country and self.year:
            lines.append(f"\nIdentified study: **{self.country} {self.year}**")

        lines.extend([
            "",
            f"### Survey Data ({len(self.data_files)})"
        ])
        for f in self.data_files:
            lines.append(f"  - {f.name}")

        if self.district_data_files:
            lines.extend([
                "",
                f"### District Data ({len(self.district_data_files)})"
            ])
            for f in self.district_data_files:
                lines.append(f"  - {f.name}")

        lines.extend([
            "",
            f"### Questionnaires ({len(self.questionnaire_files)})"
        ])
        for f in self.questionnaire_files:
            lines.append(f"  - {f.name}")

        if self.codebook_files:
            lines.extend([
                "",
                f"### Codebooks ({len(self.codebook_files)})"
            ])
            for f in self.codebook_files:
                lines.append(f"  - {f.name}")

        if self.design_report_files:
            lines.extend([
                "",
                f"### Design Reports ({len(self.design_report_files)})"
            ])
            for f in self.design_report_files:
                lines.append(f"  - {f.name}")

        if self.macro_report_files:
            lines.extend([
                "",
                f"### Macro Reports ({len(self.macro_report_files)})"
            ])
            for f in self.macro_report_files:
                lines.append(f"  - {f.name}")

        if self.other_files:
            lines.extend([
                "",
                f"### Other files ({len(self.other_files)})"
            ])
            for f in self.other_files[:10]:
                lines.append(f"  - {f.name}")
            if len(self.other_files) > 10:
                lines.append(f"  ... and {len(self.other_files) - 10} more")

        # Status
        lines.extend(["", "### Status"])
        if self.has_minimum_requirements():
            lines.append("Ready to process (data + questionnaire found)")
        else:
            lines.append("Missing required files:")
            if not self.data_files:
                lines.append("  - No data file detected")
            if not self.questionnaire_files:
                lines.append("  - No questionnaire detected")

        return "\n".join(lines)


def detect_questionnaire_language(filename: str) -> str:
    """
    Detect if a questionnaire is the English translation or native version.

    Returns:
        'english' or 'native'
    """
    name_lower = filename.lower()

    if any(kw in name_lower for kw in ENGLISH_KEYWORDS):
        return "english"

    return "native"


def parse_country_year_from_folder(folder_name: str) -> tuple[str, str]:
    """
    Parse country name and year from folder names.

    Handles various formats:
    - SouthKorea_2024 -> ("South Korea", "2024")
    - South_Korea_2024 -> ("South Korea", "2024")
    - Korea_2024 -> ("Korea", "2024")
    - NewZealand_2025 -> ("New Zealand", "2025")

    Args:
        folder_name: Name of the folder (e.g., "SouthKorea_2024")

    Returns:
        Tuple of (country_name, year). Year may be empty if not found.
    """
    # Try to extract year (4 digits at end after underscore)
    year_match = re.search(r'_(\d{4})$', folder_name)
    year = year_match.group(1) if year_match else ""

    # Remove year from name
    if year:
        name_part = folder_name[:year_match.start()]
    else:
        name_part = folder_name

    # Replace underscores with spaces
    name_part = name_part.replace('_', ' ')

    # Split camelCase (e.g., "SouthKorea" -> "South Korea")
    # Insert space before capital letters that follow lowercase
    name_part = re.sub(r'([a-z])([A-Z])', r'\1 \2', name_part)

    # Clean up extra spaces
    country = ' '.join(name_part.split())

    return (country, year)


class FileOrganizer:
    """
    Detects and organizes collaborator files into a clean folder structure.

    Creates a study folder following CSES standard structure:

    Sweden_2024/                          # Full country name
    ├── micro/
    │   ├── original_deposit/             # ALL original materials together, unchanged
    │   │   ├── original_data.dta
    │   │   ├── questionnaire.pdf
    │   │   ├── codebook.docx
    │   │   └── design_report.pdf
    │   ├── FINAL dataset/                # Cleaned/processed data
    │   │   └── cses-m6_micro_SWE_2024_*.dta
    │   ├── deposited variable list/      # Variable tracking sheet
    │   │   └── deposited variables-m6_SWE_2024_*.xlsx
    │   ├── Collaborator Questions/
    │   ├── data_checks/
    │   ├── frequencies/
    │   ├── labels/
    │   └── cses-m6_micro_SWE_2024_*.do
    ├── macro/
    ├── Election Results/
    └── .cses/                            # Agent state (hidden)
        └── state.json

    Note: Internal file prefixes use CODE_YEAR format (e.g., SWE_2024)
    """

    # Country code mapping (name -> code)
    COUNTRY_CODES = {
        "australia": "AUS", "austria": "AUT", "belgium": "BEL",
        "brazil": "BRA", "bulgaria": "BGR", "canada": "CAN",
        "chile": "CHL", "costa rica": "CRI", "croatia": "HRV",
        "czech": "CZE", "czechia": "CZE", "denmark": "DNK",
        "estonia": "EST", "finland": "FIN", "france": "FRA",
        "germany": "DEU", "great britain": "GBR", "uk": "GBR",
        "united kingdom": "GBR", "greece": "GRC", "hungary": "HUN",
        "iceland": "ISL", "ireland": "IRL", "israel": "ISR",
        "italy": "ITA", "japan": "JPN", "korea": "KOR",
        "south korea": "KOR", "latvia": "LVA", "lithuania": "LTU",
        "mexico": "MEX", "netherlands": "NLD", "new zealand": "NZL",
        "norway": "NOR", "peru": "PER", "philippines": "PHL",
        "poland": "POL", "portugal": "PRT", "romania": "ROU",
        "serbia": "SRB", "slovakia": "SVK", "slovenia": "SVN",
        "south africa": "ZAF", "spain": "ESP", "sweden": "SWE",
        "switzerland": "CHE", "taiwan": "TWN", "thailand": "THA",
        "turkey": "TUR", "united states": "USA", "usa": "USA",
        "uruguay": "URY"
    }

    # Reverse mapping (code -> full country name for folder naming)
    COUNTRY_NAMES = {
        "AUS": "Australia", "AUT": "Austria", "BEL": "Belgium",
        "BRA": "Brazil", "BGR": "Bulgaria", "CAN": "Canada",
        "CHL": "Chile", "CRI": "Costa Rica", "HRV": "Croatia",
        "CZE": "Czech Republic", "DNK": "Denmark", "EST": "Estonia",
        "FIN": "Finland", "FRA": "France", "DEU": "Germany",
        "GBR": "Great Britain", "GRC": "Greece", "HUN": "Hungary",
        "ISL": "Iceland", "IRL": "Ireland", "ISR": "Israel",
        "ITA": "Italy", "JPN": "Japan", "KOR": "Korea",
        "LVA": "Latvia", "LTU": "Lithuania", "MEX": "Mexico",
        "NLD": "Netherlands", "NZL": "New Zealand", "NOR": "Norway",
        "PER": "Peru", "PHL": "Philippines", "POL": "Poland",
        "PRT": "Portugal", "ROU": "Romania", "SRB": "Serbia",
        "SVK": "Slovakia", "SVN": "Slovenia", "ZAF": "South Africa",
        "ESP": "Spain", "SWE": "Sweden", "CHE": "Switzerland",
        "TWN": "Taiwan", "THA": "Thailand", "TUR": "Turkey",
        "USA": "United States", "URY": "Uruguay"
    }

    def __init__(self, working_dir: Path = None):
        """Initialize organizer for a directory."""
        self.working_dir = working_dir or Path.cwd()

    def find_email_folder(self) -> Path | None:
        """
        Find email folder containing deposited files.

        Returns:
            Path to email folder or None if not found
        """
        # Match: emails, Emails, E-mails, E-mail, email, etc.
        email_variants = ['emails', 'e-mails', 'e-mail', 'email']

        for item in self.working_dir.iterdir():
            if not item.is_dir():
                continue
            # Normalize: lowercase, remove hyphens/underscores
            normalized = item.name.lower().replace('-', '').replace('_', '')
            if normalized not in [v.replace('-', '') for v in email_variants]:
                continue

            # Found email folder - verify it has subfolders with data files
            subfolders = [d for d in item.iterdir() if d.is_dir()]
            if not subfolders:
                continue

            # Check if any subfolder has data files
            for subfolder in subfolders:
                if any(f.suffix.lower() in DATA_EXTENSIONS for f in subfolder.rglob("*") if f.is_file()):
                    return item

        return None

    def detect_files(self, source_dir: Path = None, recursive: bool = False) -> DetectedFiles:
        """
        Detect and classify files in the specified or working directory.

        Args:
            source_dir: Directory to scan (defaults to working directory)
            recursive: If True, scan all subdirectories recursively

        Returns:
            DetectedFiles with classified files
        """
        result = DetectedFiles()
        scan_dir = source_dir or self.working_dir

        if recursive:
            # Recursively scan all files in all subdirectories
            for item in scan_dir.rglob("*"):
                if item.name.startswith("."):
                    continue
                if item.is_dir():
                    continue

                self._classify_file(item, result)
        else:
            # Only scan immediate directory
            for item in scan_dir.iterdir():
                if item.name.startswith("."):
                    continue
                if item.is_dir():
                    continue

                self._classify_file(item, result)

        self._detect_study_info(result)
        return result

    def _classify_file(self, file_path: Path, result: DetectedFiles):
        """Classify a single file by type."""
        ext = file_path.suffix.lower()
        name_lower = file_path.name.lower()

        if ext in DATA_EXTENSIONS:
            # Check if it's district data (election results by constituency)
            if any(kw in name_lower for kw in DISTRICT_DATA_KEYWORDS):
                result.district_data_files.append(file_path)
            else:
                result.data_files.append(file_path)
            return

        if ext in DOC_EXTENSIONS:
            if any(kw in name_lower for kw in QUESTIONNAIRE_KEYWORDS):
                result.questionnaire_files.append(file_path)
            elif any(kw in name_lower for kw in CODEBOOK_KEYWORDS):
                result.codebook_files.append(file_path)
            elif any(kw in name_lower for kw in MACRO_REPORT_KEYWORDS):
                result.macro_report_files.append(file_path)
            elif any(kw in name_lower for kw in DESIGN_KEYWORDS):
                result.design_report_files.append(file_path)
            else:
                result.other_files.append(file_path)
            return

        result.other_files.append(file_path)

    def _detect_study_info(self, result: DetectedFiles):
        """Try to detect country and year from filenames."""
        all_files = (
            result.data_files +
            result.district_data_files +
            result.questionnaire_files +
            result.codebook_files +
            result.design_report_files +
            result.macro_report_files
        )

        for file_path in all_files:
            name = file_path.stem.lower()

            year_match = re.search(r'(202[0-9])', name)
            if year_match and not result.year:
                result.year = year_match.group(1)

            code_match = re.search(r'\b([A-Z]{3})\b', file_path.stem)
            if code_match:
                code = code_match.group(1)
                for country, c_code in self.COUNTRY_CODES.items():
                    if c_code == code:
                        result.country_code = code
                        result.country = country.title()
                        break

            for country, code in self.COUNTRY_CODES.items():
                if country in name:
                    result.country = country.title()
                    result.country_code = code
                    break

    def get_study_prefix(self, country: str, year: str) -> str:
        """Get the standard prefix for file naming (e.g., KOR_2024)."""
        country_code = self.COUNTRY_CODES.get(country.lower(), country[:3].upper())
        return f"{country_code}_{year}"

    def get_study_folder_name(self, country: str, year: str) -> str:
        """
        Get the folder name in CSES standard format (e.g., Sweden_2024).

        Uses full country name for folder naming, with proper capitalization.
        Falls back to the country name provided if not in mapping.

        Args:
            country: Country name (any case)
            year: Election year

        Returns:
            Folder name like "Sweden_2024"
        """
        # Get country code first
        country_code = self.COUNTRY_CODES.get(country.lower(), country[:3].upper())

        # Look up full name from code, or use provided name with title case
        full_name = self.COUNTRY_NAMES.get(country_code, country.title())

        # Replace spaces with nothing for folder name (e.g., "New Zealand" -> "NewZealand")
        # Actually CSES uses spaces in folder names, but we'll use underscores to be safe
        folder_name = full_name.replace(" ", "_")

        return f"{folder_name}_{year}"

    def create_study_folder(
        self,
        country: str,
        year: str,
        base_dir: Path = None
    ) -> Path:
        """
        Create the study folder with CSES standard structure.

        Creates:
        - micro/original_deposit/
        - micro/FINAL dataset/
        - micro/deposited variable list/
        - micro/labels/
        - micro/data_checks/
        - micro/frequencies/
        - micro/Collaborator Questions/
        - macro/
        - Election Results/
        - .cses/

        Args:
            country: Country name
            year: Election year
            base_dir: Base directory (default: working_dir)

        Returns:
            Path to created study folder
        """
        if base_dir is None:
            base_dir = self.working_dir

        folder_name = self.get_study_folder_name(country, year)
        study_dir = base_dir / folder_name

        study_dir.mkdir(parents=True, exist_ok=True)

        # Create CSES standard folder structure
        # micro/ subfolders
        (study_dir / "micro" / "original_deposit").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "FINAL dataset").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "deposited variable list").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "labels").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "data_checks").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "frequencies").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "Collaborator Questions").mkdir(parents=True, exist_ok=True)

        # macro/
        (study_dir / "macro").mkdir(exist_ok=True)

        # Election Results/
        (study_dir / "Election Results").mkdir(exist_ok=True)

        # .cses/ for agent state
        (study_dir / ".cses").mkdir(exist_ok=True)

        logger.info(f"Created study folder: {study_dir}")
        return study_dir

    def create_study_structure(self, study_dir: Path):
        """
        Create CSES folder structure at the specified directory (root level).

        This is used when the country folder already exists (e.g., from email deposit)
        and we just need to create the processing folders alongside the email folder.

        Creates:
        - micro/
        - micro/FINAL dataset/
        - micro/deposited variable list/
        - micro/Collaborator Questions/
        - macro/
        - Election Results/
        - .cses/

        Note: Does NOT create micro/original_deposit/ since the email folder
        serves as the raw backup.

        Args:
            study_dir: Directory to create structure in
        """
        # micro/ and subfolders
        (study_dir / "micro").mkdir(exist_ok=True)
        (study_dir / "micro" / "FINAL dataset").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "deposited variable list").mkdir(parents=True, exist_ok=True)
        (study_dir / "micro" / "Collaborator Questions").mkdir(parents=True, exist_ok=True)

        # macro/
        (study_dir / "macro").mkdir(exist_ok=True)

        # Election Results/
        (study_dir / "Election Results").mkdir(exist_ok=True)

        # .cses/ for agent state
        (study_dir / ".cses").mkdir(exist_ok=True)

        logger.info(f"Created study structure in: {study_dir}")

    def _copy_as_pdf(self, src: Path, dst: Path):
        """
        Copy a document file, converting to PDF if necessary.

        Handles: .pdf (direct copy), .docx/.doc, .txt, .rtf
        Preserves formatting as much as possible (headings, bold, lists).
        If conversion fails, copies the original file alongside.

        Args:
            src: Source file path
            dst: Destination path (should end in .pdf)
        """
        src_ext = src.suffix.lower()

        if src_ext == '.pdf':
            # Direct copy
            shutil.copy2(src, dst)
            print(f"    [OK] Copied PDF")
            return

        # Try to convert .docx to PDF with formatting
        if src_ext in ['.docx', '.doc']:
            try:
                from docx import Document
                from docx.shared import Pt
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib import colors

                # Read docx
                doc = Document(src)

                # Create PDF with platypus for better formatting
                pdf_doc = SimpleDocTemplate(str(dst), pagesize=letter,
                                           leftMargin=inch, rightMargin=inch,
                                           topMargin=inch, bottomMargin=inch)

                styles = getSampleStyleSheet()
                # Add custom styles
                styles.add(ParagraphStyle(name='Heading1Custom',
                                         parent=styles['Heading1'],
                                         fontSize=14, spaceAfter=12))
                styles.add(ParagraphStyle(name='Heading2Custom',
                                         parent=styles['Heading2'],
                                         fontSize=12, spaceAfter=10))
                styles.add(ParagraphStyle(name='BodyCustom',
                                         parent=styles['Normal'],
                                         fontSize=10, spaceAfter=6))

                story = []

                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        story.append(Spacer(1, 6))
                        continue

                    # Escape special characters for reportlab
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                    # Detect style based on paragraph properties
                    style_name = 'BodyCustom'
                    if para.style and para.style.name:
                        style_lower = para.style.name.lower()
                        if 'heading 1' in style_lower or 'title' in style_lower:
                            style_name = 'Heading1Custom'
                        elif 'heading 2' in style_lower:
                            style_name = 'Heading2Custom'
                        elif 'heading' in style_lower:
                            style_name = 'Heading2Custom'

                    # Check for bold runs (make whole para bold if first run is bold)
                    if para.runs and para.runs[0].bold:
                        text = f"<b>{text}</b>"

                    story.append(Paragraph(text, styles[style_name]))

                # Handle tables
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            row_data.append(Paragraph(cell_text, styles['BodyCustom']))
                        table_data.append(row_data)

                    if table_data:
                        t = Table(table_data)
                        t.setStyle(TableStyle([
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ]))
                        story.append(Spacer(1, 12))
                        story.append(t)
                        story.append(Spacer(1, 12))

                pdf_doc.build(story)
                print(f"    [OK] Converted to PDF")
                return

            except ImportError as e:
                print(f"    [!] Cannot convert to PDF - missing dependency: {e}")
            except Exception as e:
                print(f"    [!] Failed to convert to PDF: {e}")

        # Try to convert .txt to PDF
        if src_ext == '.txt':
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch

                text_content = src.read_text(encoding='utf-8', errors='replace')

                pdf_doc = SimpleDocTemplate(str(dst), pagesize=letter,
                                           leftMargin=inch, rightMargin=inch,
                                           topMargin=inch, bottomMargin=inch)
                styles = getSampleStyleSheet()
                story = []

                for line in text_content.split('\n'):
                    line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    if line.strip():
                        story.append(Paragraph(line, styles['Normal']))
                    else:
                        story.append(Spacer(1, 6))

                pdf_doc.build(story)
                print(f"    [OK] Converted to PDF")
                return

            except Exception as e:
                print(f"    [!] Failed to convert txt to PDF: {e}")

        # Fallback: copy original file (keep original extension)
        fallback_dst = dst.with_suffix(src_ext)
        shutil.copy2(src, fallback_dst)
        print(f"    [!] Could not convert to PDF, kept original format: {fallback_dst.name}")

    def _translate_questionnaire_to_english(self, src: Path, dst: Path):
        """
        Translate a questionnaire to English using LLM and save as PDF.

        Adds a disclaimer that the document was AI-translated.

        Args:
            src: Source file path (native language questionnaire)
            dst: Destination path (should end in .pdf)
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch

        # Extract text from source document
        src_ext = src.suffix.lower()
        source_text = ""

        if src_ext == '.pdf':
            try:
                from pypdf import PdfReader
                reader = PdfReader(src)
                for page in reader.pages:
                    source_text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"Could not read PDF: {e}")
        elif src_ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(src)
                source_text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                logger.warning(f"Could not read docx: {e}")
        elif src_ext == '.txt':
            source_text = src.read_text(encoding='utf-8', errors='replace')

        if not source_text.strip():
            logger.warning(f"Could not extract text from {src.name}, copying as-is")
            self._copy_as_pdf(src, dst)
            return

        # Translate using LLM
        try:
            from litellm import completion
            import os

            model = os.getenv("LLM_MODEL_PREPROCESS", "openai/gpt-4o-mini")

            print(f"Translating questionnaire to English...")
            response = completion(
                model=model,
                max_tokens=8192,
                timeout=300,
                messages=[{
                    "role": "user",
                    "content": f"""Translate this questionnaire to English. Preserve the structure and formatting.

QUESTIONNAIRE TEXT:
{source_text[:15000]}

Translate to English, keeping question numbers and response options intact."""
                }]
            )

            translated_text = response.choices[0].message.content.strip()

        except Exception as e:
            logger.error(f"Translation failed: {e}")
            translated_text = f"[TRANSLATION FAILED: {e}]\n\n{source_text}"

        # Create PDF with disclaimer
        c = canvas.Canvas(str(dst), pagesize=letter)
        width, height = letter
        y_position = height - inch

        # Add disclaimer at top
        c.setFont("Helvetica-Bold", 12)
        c.drawString(inch, y_position, "DISCLAIMER: This document was translated by AI.")
        y_position -= 16
        c.setFont("Helvetica-Oblique", 10)
        c.drawString(inch, y_position, "Original document should be consulted for accuracy.")
        y_position -= 24
        c.drawString(inch, y_position, "-" * 70)
        y_position -= 20

        # Add translated content
        c.setFont("Helvetica", 10)
        for line in translated_text.split('\n'):
            line = line.strip()
            if not line:
                y_position -= 10
                continue

            # Wrap long lines
            words = line.split()
            current_line = ""
            for word in words:
                test_line = f"{current_line} {word}".strip()
                if len(test_line) > 85:
                    c.drawString(inch, y_position, current_line)
                    y_position -= 12
                    current_line = word
                else:
                    current_line = test_line

            if current_line:
                c.drawString(inch, y_position, current_line)
                y_position -= 12

            # New page if needed
            if y_position < inch:
                c.showPage()
                c.setFont("Helvetica", 10)
                y_position = height - inch

        c.save()
        logger.info(f"Created AI-translated English questionnaire: {dst.name}")

    def copy_files_with_standard_names(
        self,
        detected: DetectedFiles,
        source_dir: Path,
        study_dir: Path,
        country_code: str,
        year: str
    ) -> dict:
        """
        Copy detected files to micro/ with standardized names.

        Naming: {CODE}_{YEAR}_{type}.{ext}
        Examples: KOR_2024_data.dta, KOR_2024_questionnaire.pdf

        Args:
            detected: DetectedFiles result
            source_dir: Where the original files are (e.g., emails/20250303/)
            study_dir: Target study directory (root of country folder)
            country_code: 3-letter country code (e.g., KOR)
            year: Election year (e.g., 2024)

        Returns:
            Dict mapping file type to path of standardized copy
        """
        mapping = {}
        micro_dir = study_dir / "micro"
        prefix = f"{country_code}_{year}"

        # Copy data file(s)
        if detected.data_files:
            src = detected.data_files[0]
            dst = micro_dir / f"{prefix}_data{src.suffix}"
            print(f"  Copying data file: {src.name}")
            shutil.copy2(src, dst)
            mapping["data"] = str(dst)

            # If multiple data files, copy with numeric suffixes
            for i, src in enumerate(detected.data_files[1:], start=2):
                dst = micro_dir / f"{prefix}_data_{i}{src.suffix}"
                print(f"  Copying data file: {src.name}")
                shutil.copy2(src, dst)
                mapping[f"data_{i}"] = str(dst)

        # Copy questionnaires - one English, one native (both as PDF)
        if detected.questionnaire_files:
            english_questionnaires = []
            native_questionnaires = []

            for src in detected.questionnaire_files:
                lang = detect_questionnaire_language(src.name)
                if lang == "english":
                    english_questionnaires.append(src)
                else:
                    native_questionnaires.append(src)

            # Select best native questionnaire (prefer PDF) - always copy native first
            native_src = None
            if native_questionnaires:
                native_questionnaires.sort(key=lambda f: (f.suffix.lower() != '.pdf', f.name))
                native_src = native_questionnaires[0]
                dst = micro_dir / f"{prefix}_questionnaire_native.pdf"
                print(f"  Converting native questionnaire to PDF: {native_src.name}")
                self._copy_as_pdf(native_src, dst)
                mapping["questionnaire_native"] = str(dst)

            # Select best English questionnaire (prefer PDF)
            if english_questionnaires:
                english_questionnaires.sort(key=lambda f: (f.suffix.lower() != '.pdf', f.name))
                src = english_questionnaires[0]
                dst = micro_dir / f"{prefix}_questionnaire_english.pdf"
                print(f"  Converting English questionnaire to PDF: {src.name}")
                self._copy_as_pdf(src, dst)
                mapping["questionnaire_english"] = str(dst)
            elif native_src:
                # No English questionnaire - translate the native one
                print("  No English questionnaire found - translating native version...")
                dst = micro_dir / f"{prefix}_questionnaire_english.pdf"
                self._translate_questionnaire_to_english(native_src, dst)
                mapping["questionnaire_english"] = str(dst)
                mapping["questionnaire_english_translated"] = True

        # Copy codebook (convert to PDF)
        if detected.codebook_files:
            src = detected.codebook_files[0]
            dst = micro_dir / f"{prefix}_codebook.pdf"
            print(f"  Converting codebook to PDF: {src.name}")
            self._copy_as_pdf(src, dst)
            mapping["codebook"] = str(dst)

        # Copy design report (convert to PDF)
        if detected.design_report_files:
            src = detected.design_report_files[0]
            dst = micro_dir / f"{prefix}_design_report.pdf"
            print(f"  Converting design report to PDF: {src.name}")
            self._copy_as_pdf(src, dst)
            mapping["design_report"] = str(dst)

        # Copy macro reports (convert to PDF)
        if detected.macro_report_files:
            src = detected.macro_report_files[0]
            dst = micro_dir / f"{prefix}_macro_report.pdf"
            print(f"  Converting macro report to PDF: {src.name}")
            self._copy_as_pdf(src, dst)
            mapping["macro_report"] = str(dst)

        # Copy district data
        if detected.district_data_files:
            src = detected.district_data_files[0]
            dst = micro_dir / f"{prefix}_district_data{src.suffix}"
            print(f"  Copying district data: {src.name}")
            shutil.copy2(src, dst)
            mapping["district_data"] = str(dst)

        return mapping

    def organize_files(
        self,
        detected: DetectedFiles,
        study_dir: Path,
        prefix: str
    ) -> dict:
        """
        Organize files following CSES standard structure.

        Copies ALL original files to micro/original_deposit/ with their
        original names preserved (unchanged). No working copies are created
        in the root folder - all processing works from original_deposit.

        Args:
            detected: DetectedFiles result
            study_dir: Target study directory
            prefix: File prefix (e.g., KOR_2024) - used for tracking

        Returns:
            Dict mapping original paths to paths in original_deposit
        """
        mapping = {}

        # Ensure micro/original_deposit exists
        original_deposit = study_dir / "micro" / "original_deposit"
        original_deposit.mkdir(parents=True, exist_ok=True)

        # Collect all source files for original deposit
        all_source_files = (
            detected.data_files +
            detected.district_data_files +
            detected.questionnaire_files +
            detected.codebook_files +
            detected.design_report_files +
            detected.macro_report_files
        )

        # Copy all originals to micro/original_deposit/ (keep original names)
        for src in all_source_files:
            dst = original_deposit / src.name
            if not dst.exists():
                shutil.copy2(src, dst)
                mapping[str(src)] = str(dst)
                logger.info(f"Preserved original: {src.name}")

        # Note: We no longer create renamed working copies in the root folder.
        # All processing now works directly from files in micro/original_deposit/

        return mapping

    def initialize_study(
        self,
        country: str = None,
        year: str = None,
        copy_files: bool = True
    ) -> Tuple[Path, DetectedFiles]:
        """
        Full initialization: detect files, create folder, organize with proper names.

        Returns:
            Tuple of (study_dir, detected_files)
        """
        detected = self.detect_files()

        country = country or detected.country or "Unknown"
        year = year or detected.year or "0000"

        logger.info(f"Initializing study: {country} {year}")

        study_dir = self.create_study_folder(country, year)
        prefix = self.get_study_prefix(country, year)

        if detected.data_files or detected.questionnaire_files:
            self.organize_files(detected, study_dir, prefix)

        return study_dir, detected

    def migrate_old_structure(self, study_dir: Path) -> dict:
        """
        Migrate an old folder structure to the new CSES standard.

        Performs:
        1. Move original_deposit/ -> micro/original_deposit/ (if at root)
        2. Create missing standard folders
        4. Returns info about what was changed

        Args:
            study_dir: The study directory to migrate

        Returns:
            Dict with migration results: moved_files, renamed_folders, created_folders
        """
        results = {
            "moved_files": [],
            "renamed_folders": [],
            "created_folders": [],
            "errors": []
        }

        # 1. Move original_deposit/ from root to micro/original_deposit/
        old_deposit = study_dir / "original_deposit"
        new_deposit = study_dir / "micro" / "original_deposit"

        if old_deposit.exists() and old_deposit.is_dir():
            # Create micro/ if needed
            (study_dir / "micro").mkdir(exist_ok=True)

            if new_deposit.exists():
                # Merge files from old to new
                for item in old_deposit.iterdir():
                    dst = new_deposit / item.name
                    if not dst.exists():
                        shutil.move(str(item), str(dst))
                        results["moved_files"].append(f"{item.name} -> micro/original_deposit/")
                        logger.info(f"Migrated: {item.name} -> micro/original_deposit/")

                # Remove empty old folder
                try:
                    old_deposit.rmdir()
                    results["renamed_folders"].append("original_deposit/ removed (merged into micro/)")
                except OSError:
                    results["errors"].append("Could not remove old original_deposit/ (not empty)")
            else:
                # Move entire folder
                shutil.move(str(old_deposit), str(new_deposit))
                results["renamed_folders"].append("original_deposit/ -> micro/original_deposit/")
                logger.info("Migrated: original_deposit/ -> micro/original_deposit/")

        # 2. Create missing standard folders
        standard_folders = [
            "micro/original_deposit",
            "micro/FINAL dataset",
            "micro/deposited variable list",
            "micro/labels",
            "micro/data_checks",
            "micro/frequencies",
            "micro/Collaborator Questions",
            "macro",
            "Election Results",
        ]

        for folder in standard_folders:
            folder_path = study_dir / folder
            if not folder_path.exists():
                folder_path.mkdir(parents=True, exist_ok=True)
                results["created_folders"].append(folder)
                logger.info(f"Created: {folder}/")

        return results


    def cleanup_source_files(
        self,
        source_dir: Path,
        study_dir: Path,
        dry_run: bool = False
    ) -> dict:
        """
        Remove source files from the root directory after they've been copied
        to micro/original_deposit/.

        Only removes files that:
        1. Exist in micro/original_deposit/ with the same name
        2. Have the same file size (safety check)

        Args:
            source_dir: Directory containing original source files
            study_dir: Study directory with micro/original_deposit/
            dry_run: If True, only report what would be deleted

        Returns:
            Dict with cleanup results: deleted, skipped, errors
        """
        results = {
            "deleted": [],
            "skipped": [],
            "errors": []
        }

        original_deposit = study_dir / "micro" / "original_deposit"
        if not original_deposit.exists():
            results["errors"].append("micro/original_deposit/ does not exist")
            return results

        # Get all files in original_deposit
        deposited_files = {f.name: f for f in original_deposit.iterdir() if f.is_file()}

        # Check source directory for matching files
        for src_file in source_dir.iterdir():
            if src_file.is_dir():
                continue
            if src_file.name.startswith("."):
                continue

            # Skip if not in original_deposit
            if src_file.name not in deposited_files:
                continue

            dst_file = deposited_files[src_file.name]

            # Safety check: verify file sizes match
            if src_file.stat().st_size != dst_file.stat().st_size:
                results["skipped"].append(
                    f"{src_file.name} (size mismatch - may have been modified)"
                )
                continue

            # Delete the source file
            if dry_run:
                results["deleted"].append(f"{src_file.name} (would delete)")
            else:
                try:
                    src_file.unlink()
                    results["deleted"].append(src_file.name)
                    logger.info(f"Cleaned up: {src_file.name}")
                except Exception as e:
                    results["errors"].append(f"{src_file.name}: {e}")

        return results


def detect_and_summarize(directory: Path = None) -> str:
    """Quick detection and summary for a directory."""
    organizer = FileOrganizer(directory)
    detected = organizer.detect_files()
    return detected.summary()
